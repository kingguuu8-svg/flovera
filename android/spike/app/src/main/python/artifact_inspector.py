import csv
import html.parser
import json
import os
import re
import struct
import traceback


TEXT_LIMIT = 4000


def inspect_artifact(workspace_root, path, max_text_chars=TEXT_LIMIT):
    root = os.path.realpath(workspace_root)
    target = _resolve(root, path)
    size = os.path.getsize(target)
    extension = os.path.splitext(target)[1].lower().lstrip(".")
    try:
        if extension == "json":
            result = _inspect_json(target, max_text_chars)
        elif extension in ("html", "htm"):
            result = _inspect_html(root, target, max_text_chars)
        elif extension == "docx":
            result = _inspect_docx(target, max_text_chars)
        elif extension == "xlsx":
            result = _inspect_xlsx(target)
        elif extension == "pdf":
            result = _inspect_pdf(target, max_text_chars)
        elif extension in ("png", "jpg", "jpeg", "webp"):
            result = _inspect_image(target, extension)
        else:
            result = _inspect_text_like(target, max_text_chars)
        result["path"] = os.path.relpath(target, root).replace(os.sep, "/")
        result["sizeBytes"] = size
        return json.dumps(result, ensure_ascii=False, separators=(",", ":"))
    except Exception as exc:
        return json.dumps(
            {
                "status": "error",
                "path": os.path.relpath(target, root).replace(os.sep, "/"),
                "format": extension or "unknown",
                "sizeBytes": size,
                "summary": "Artifact inspection failed.",
                "problems": [str(exc)],
                "structuredFacts": {"traceback": traceback.format_exc(limit=5)},
            },
            ensure_ascii=False,
            separators=(",", ":"),
        )


def _resolve(root, path):
    target = os.path.realpath(os.path.join(root, path))
    if target != root and not target.startswith(root + os.sep):
        raise PermissionError("Path escapes workspace: " + path)
    if not os.path.isfile(target):
        raise FileNotFoundError(path)
    return target


def _inspect_json(path, max_text_chars):
    with open(path, "r", encoding="utf-8") as handle:
        parsed = json.load(handle)
    facts = {"topLevelType": type(parsed).__name__}
    if isinstance(parsed, dict):
        facts["topLevelKeys"] = list(parsed.keys())[:50]
        facts["keyCount"] = len(parsed)
    elif isinstance(parsed, list):
        facts["itemCount"] = len(parsed)
        facts["firstItemType"] = type(parsed[0]).__name__ if parsed else "empty"
    return {
        "status": "ok",
        "format": "json",
        "summary": "Valid JSON " + facts["topLevelType"],
        "problems": [],
        "structuredFacts": facts,
        "textPreview": _text_preview(path, max_text_chars),
    }


class _HtmlFactsParser(html.parser.HTMLParser):
    def __init__(self):
        super().__init__()
        self.title = ""
        self._in_title = False
        self.links = []
        self.scripts = []
        self.images = []

    def handle_starttag(self, tag, attrs):
        values = dict(attrs)
        if tag == "title":
            self._in_title = True
        if tag == "link" and values.get("href"):
            self.links.append(values["href"])
        if tag == "script" and values.get("src"):
            self.scripts.append(values["src"])
        if tag == "img" and values.get("src"):
            self.images.append(values["src"])

    def handle_endtag(self, tag):
        if tag == "title":
            self._in_title = False

    def handle_data(self, data):
        if self._in_title:
            self.title += data.strip()


def _inspect_html(root, path, max_text_chars):
    text = _read_text(path, max_text_chars * 4)
    parser = _HtmlFactsParser()
    parser.feed(text)
    refs = parser.links + parser.scripts + parser.images
    missing = []
    base = os.path.dirname(path)
    for ref in refs:
        if re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", ref) or ref.startswith("#"):
            continue
        candidate = os.path.realpath(os.path.join(base, ref.split("#", 1)[0].split("?", 1)[0]))
        if candidate.startswith(root + os.sep) and not os.path.exists(candidate):
            missing.append(os.path.relpath(candidate, root).replace(os.sep, "/"))
    return {
        "status": "warn" if missing else "ok",
        "format": "html",
        "summary": "HTML title=" + (parser.title or "(none)"),
        "problems": ["Missing local resource: " + item for item in missing[:20]],
        "structuredFacts": {
            "title": parser.title,
            "stylesheetCount": len(parser.links),
            "scriptCount": len(parser.scripts),
            "imageCount": len(parser.images),
            "missingResourceCount": len(missing),
        },
        "textPreview": text[:max_text_chars],
    }


def _inspect_docx(path, max_text_chars):
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return {
        "status": "ok",
        "format": "docx",
        "summary": f"DOCX paragraphs={len(paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)}",
        "problems": [],
        "structuredFacts": {
            "paragraphCount": len(paragraphs),
            "tableCount": len(doc.tables),
            "inlineImageCount": len(doc.inline_shapes),
            "firstParagraphs": paragraphs[:10],
        },
        "textPreview": text[:max_text_chars],
    }


def _inspect_xlsx(path):
    import openpyxl

    workbook = openpyxl.load_workbook(path, data_only=False, read_only=True)
    sheets = []
    formula_count = 0
    non_empty_count = 0
    for sheet in workbook.worksheets:
        sheet_non_empty = 0
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    non_empty_count += 1
                    sheet_non_empty += 1
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        formula_count += 1
        sheets.append({"name": sheet.title, "dimension": sheet.calculate_dimension(), "nonEmptyCells": sheet_non_empty})
    workbook.close()
    return {
        "status": "ok",
        "format": "xlsx",
        "summary": f"XLSX sheets={len(sheets)} cells={non_empty_count} formulas={formula_count}",
        "problems": [],
        "structuredFacts": {"sheets": sheets, "nonEmptyCellCount": non_empty_count, "formulaCount": formula_count},
    }


def _inspect_pdf(path, max_text_chars):
    from pypdf import PdfReader

    reader = PdfReader(path)
    text_parts = []
    for page in reader.pages[:3]:
        text_parts.append(page.extract_text() or "")
    text = "\n".join(part for part in text_parts if part)
    metadata = {str(key): str(value) for key, value in (reader.metadata or {}).items()}
    return {
        "status": "ok",
        "format": "pdf",
        "summary": f"PDF pages={len(reader.pages)} textChars={len(text)}",
        "problems": [],
        "structuredFacts": {"pageCount": len(reader.pages), "metadata": metadata, "encrypted": reader.is_encrypted},
        "textPreview": text[:max_text_chars],
    }


def _inspect_image(path, extension):
    width = None
    height = None
    with open(path, "rb") as handle:
        header = handle.read(64)
        if extension == "png" and header.startswith(b"\x89PNG\r\n\x1a\n"):
            width, height = struct.unpack(">II", header[16:24])
        elif extension in ("jpg", "jpeg") and header.startswith(b"\xff\xd8"):
            width, height = _jpeg_size(path)
        elif extension == "webp" and header[0:4] == b"RIFF" and header[8:12] == b"WEBP":
            width, height = _webp_size(header)
    problems = [] if width and height else ["Could not determine image dimensions."]
    return {
        "status": "ok" if not problems else "warn",
        "format": extension,
        "summary": f"{extension.upper()} image {width or '?'}x{height or '?'}",
        "problems": problems,
        "structuredFacts": {"width": width, "height": height},
    }


def _jpeg_size(path):
    with open(path, "rb") as handle:
        handle.read(2)
        while True:
            marker_prefix = handle.read(1)
            if marker_prefix != b"\xff":
                return None, None
            marker = handle.read(1)
            while marker == b"\xff":
                marker = handle.read(1)
            length_data = handle.read(2)
            if len(length_data) != 2:
                return None, None
            length = struct.unpack(">H", length_data)[0]
            if marker in [bytes([m]) for m in range(0xC0, 0xC4)]:
                handle.read(1)
                height, width = struct.unpack(">HH", handle.read(4))
                return width, height
            handle.seek(length - 2, os.SEEK_CUR)


def _webp_size(header):
    if header[12:16] == b"VP8X":
        width = 1 + int.from_bytes(header[24:27], "little")
        height = 1 + int.from_bytes(header[27:30], "little")
        return width, height
    return None, None


def _inspect_text_like(path, max_text_chars):
    text = _read_text(path, max_text_chars)
    line_count = text.count("\n") + (1 if text else 0)
    return {
        "status": "ok",
        "format": "text",
        "summary": f"Text file previewChars={len(text)} lines={line_count}",
        "problems": [],
        "structuredFacts": {"previewLineCount": line_count},
        "textPreview": text,
    }


def _text_preview(path, max_text_chars):
    return _read_text(path, max_text_chars)


def _read_text(path, max_chars):
    with open(path, "r", encoding="utf-8") as handle:
        return handle.read(max_chars)
